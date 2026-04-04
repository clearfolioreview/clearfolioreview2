#!/usr/bin/env python3
"""
╔══════════════════════════════════════════════════════════════════════════════╗
║      CLEARFOLIO REVIEW v3 — Production-Grade AI Contract Audit System       ║
║      Privacy-First | Local-Only | Multi-Model | Judge Consensus              ║
╠══════════════════════════════════════════════════════════════════════════════╣
║  Fixes applied vs v2:                                                        ║
║  #01  _fatal/_warn/_status moved before first use                            ║
║  #02  RuleEngine dedupe scoped by (rule_id, clause_index)                    ║
║  #03  Clause index re-assigned post-build; fallback indices consistent       ║
║  #04  AI clause matching uses embedded [idx] prefix — exact, not fuzzy       ║
║  #05  OllamaClient.generate retries 3× with exponential back-off             ║
║  #06  Model resolution handles :tag suffixes in both directions              ║
║  #07  DocumentLoader caps full_text at MAX_DOC_CHARS (200 k chars)           ║
║  #08  All 30 rule regexes precompiled at class definition                    ║
║  #09  Table cell extraction deduplicates with seen-set                       ║
║  #10  parse_model_response uses re.split(maxsplit) — extra pipes safe        ║
║  #11  Ollama-skipped banner is explicit; ai_skipped flag on report           ║
║  #12  AuditScorer accumulates penalties separately; clamping logic clean     ║
║  #13  logging module used throughout (--verbose enables DEBUG)               ║
║  #14  --output json|txt flag writes structured report file                   ║
║  #15  Windows ANSI colours initialised via ctypes                            ║
║  #16  Heading regex handles multi-digit levels (Heading 10, 11 …)            ║
║  #17  ThreadPool uses per-future .result(timeout=) instead of iterator TO   ║
║  #18  Judge fallback key includes issue prefix — no hash collisions          ║
║  #19  Simple SHA-256 prompt cache avoids redundant Ollama calls              ║
║  #20  Deep mode adds polite inter-pass sleep; progress shown correctly       ║
║  #21  Explicit model validation — fails fast if no model resolvable          ║
║  #22  DocumentLoader handles >200k chars with truncation warning             ║
║  #23  parse_confidence handles empty string / zero-division edge cases       ║
║  #24  QuickAnalyser wraps AIFinding natively — no dict/dataclass mismatch    ║
║  #25  Section on architectural notes added to --help / --info flag           ║
╚══════════════════════════════════════════════════════════════════════════════╝

Usage:
    python clearfolio_review.py <contract.docx> [options]

Options:
    --mode   quick|deep          quick = single model (default), deep = 3 + judge
    --output json|txt            write report to file alongside input
    --verbose                    enable DEBUG logging
    --info                       print architecture notes and exit

Dependencies:
    pip install python-docx requests

Ollama setup:
    ollama pull llama3            # minimum — used as quick + risk model
    ollama pull mistral           # compliance model (deep mode)
    ollama pull phi3              # completeness model (deep mode)
    ollama pull deepseek-r1       # preferred judge (falls back to llama3)
"""

# ─────────────────────────────────────────────────────────────────────────────
# STDLIB IMPORTS
# ─────────────────────────────────────────────────────────────────────────────
import sys
import os
import re
import json
import hashlib
import logging
import textwrap
import time
import argparse
import ctypes
import threading
from datetime import datetime
from concurrent.futures import ThreadPoolExecutor
from dataclasses import dataclass, field, asdict
from typing import Optional

# Third-party
import requests
from docx import Document


# ─────────────────────────────────────────────────────────────────────────────
# WINDOWS ANSI SUPPORT  [Fix #15]
# ─────────────────────────────────────────────────────────────────────────────
def _enable_windows_ansi() -> None:
    if sys.platform != "win32":
        return

    try:
        kernel32 = ctypes.windll.kernel32  # type: ignore[attr-defined]

        STD_OUTPUT_HANDLE = -11
        ENABLE_VIRTUAL_TERMINAL_PROCESSING = 0x0004

        handle = kernel32.GetStdHandle(STD_OUTPUT_HANDLE)
        mode = kernel32.GetConsoleMode(handle)
        kernel32.SetConsoleMode(
            handle, mode | ENABLE_VIRTUAL_TERMINAL_PROCESSING
        )

    except Exception:
        pass

    _enable_windows_ansi()

# ─────────────────────────────────────────────────────────────────────────────
# LOGGING  [Fix #13]
# ─────────────────────────────────────────────────────────────────────────────
log = logging.getLogger("clearfolio")


def _configure_logging(verbose: bool) -> None:
    level = logging.DEBUG if verbose else logging.WARNING
    logging.basicConfig(
        format="%(asctime)s  %(levelname)-8s  %(name)s  %(message)s",
        datefmt="%H:%M:%S",
        level=level,
        stream=sys.stderr,
    )
    log.setLevel(level)


# ─────────────────────────────────────────────────────────────────────────────
# COLOUR HELPERS
# ─────────────────────────────────────────────────────────────────────────────
class C:
    RESET = "\033[0m"
    BOLD = "\033[1m"
    DIM = "\033[2m"
    RED = "\033[91m"
    YELLOW = "\033[93m"
    GREEN = "\033[92m"
    CYAN = "\033[96m"
    MAGENTA = "\033[95m"
    WHITE = "\033[97m"
    BLUE = "\033[94m"
    ORANGE = "\033[38;5;208m"
    PURPLE = "\033[38;5;141m"


def col(text: object, *codes: str) -> str:
    return "".join(codes) + str(text) + C.RESET


def hr(char: str = "─", colour: str = C.DIM) -> str:
    return col(char * TERMINAL_WIDTH, colour)


# ─────────────────────────────────────────────────────────────────────────────
# UTILITIES  [Fix #01 — defined before any use]
# ─────────────────────────────────────────────────────────────────────────────
def _fatal(msg: str) -> None:
    log.critical(msg)
    print(col(f"\n  ✗ FATAL: {msg}\n", C.RED, C.BOLD), file=sys.stderr)
    sys.exit(1)


def _warn(msg: str) -> None:
    log.warning(msg)
    print(col(f"  ⚠  {msg}", C.YELLOW))


def _status(msg: str) -> None:
    log.info(msg)
    print(col(f"  ▶  {msg}", C.CYAN))


def _section(msg: str) -> None:
    log.debug(f"--- SECTION: {msg}")
    print()
    print(col(f"  {'─' * 4}  {msg}  {'─' * 4}", C.BOLD, C.WHITE))
    print()


def _banner_skipped(reason: str) -> None:
    """Explicit visible banner when AI analysis is skipped  [Fix #11]"""
    print()
    print(hr("─"))
    print(col(f"  ⊘  AI ANALYSIS SKIPPED", C.YELLOW, C.BOLD))
    print(col(f"     Reason: {reason}", C.DIM))
    print(col(f"     Rule-engine findings are still complete.", C.DIM))
    print(hr("─"))
    print()


# ─────────────────────────────────────────────────────────────────────────────
# CONFIGURATION
# ─────────────────────────────────────────────────────────────────────────────
OLLAMA_BASE_URL = os.getenv("CF_OLLAMA_URL", "http://localhost:11434")
OLLAMA_TIMEOUT = int(os.getenv("CF_TIMEOUT", "180"))  # per request
OLLAMA_RETRIES = int(os.getenv("CF_RETRIES", "3"))  # [Fix #05]
OLLAMA_RETRY_DELAY = float(os.getenv("CF_RETRY_DELAY", "2.0"))

TERMINAL_WIDTH = 92
MAX_DOC_CHARS = 500_000  # [Fix #07]  ~40 k words cap

DEEP_MAX_CLAUSES_PASS1 = None
DEEP_MAX_CLAUSES_PASS2 = None
QUICK_MAX_CLAUSES = 25
MAX_CHARS_PER_CLAUSE = 1200

INTER_PASS_SLEEP = float(os.getenv("CF_PASS_SLEEP", "1.5"))  # [Fix #20]

# ── Model env-var config ─────────────────────────────────────────────────────
MODEL_QUICK = os.getenv("CF_MODEL_QUICK", "llama3")
MODEL_A = os.getenv("CF_MODEL_A", "llama3")  # risk
MODEL_B = os.getenv("CF_MODEL_B", "mistral")  # compliance
MODEL_C = os.getenv("CF_MODEL_C", "phi3")  # completeness
JUDGE_MODEL = os.getenv("CF_JUDGE_MODEL", "deepseek-r1")  # judge


# ─────────────────────────────────────────────────────────────────────────────
# DATA STRUCTURES
# ─────────────────────────────────────────────────────────────────────────────
@dataclass
class AIFinding:
    clause_label: str
    risk: str  # CRITICAL / HIGH / MEDIUM / LOW
    issue: str
    suggestion: str
    confidence: int  # 0–100
    clause_index: int = -1
    source_models: list = field(default_factory=list)
    judge_validated: bool = False
    judge_note: str = ""
    pass_number: int = 1


@dataclass
class ModelResult:
    model_name: str
    role: str
    raw_output: str
    findings: list
    elapsed_sec: float = 0.0
    error: str = ""
    pass_number: int = 1


# ─────────────────────────────────────────────────────────────────────────────
# 1. DOCUMENT LOADER
# ─────────────────────────────────────────────────────────────────────────────
class DocumentLoader:
    SCAN_KEYWORDS = [
        "scanned", "scan only", "image-based", "ocr required",
        "this document is an image",
    ]

    def __init__(self, filepath: str):
        self.filepath = filepath
        self.doc = None
        self.full_text = ""
        self.paragraphs: list[dict] = []
        self.truncated = False

    def load(self) -> "DocumentLoader":
        self._validate_path()
        self._open_docx()
        self._extract_paragraphs()
        self._check_not_scanned()
        self._build_full_text()
        return self

    # ── private ──────────────────────────────────────────────────────────────
    def _validate_path(self) -> None:
        if not os.path.exists(self.filepath):
            _fatal(f"File not found: {self.filepath}")
        if not self.filepath.lower().endswith(".docx"):
            _fatal("Only .docx files are accepted.")

    def _open_docx(self) -> None:
        try:
            self.doc = Document(self.filepath)
        except Exception as exc:
            _fatal(f"Cannot open .docx: {exc}")

    def _extract_paragraphs(self) -> None:
        """Extract paragraphs + table cells with deduplication.  [Fix #09]"""
        seen_texts: set[str] = set()

        def _add(text: str, style: str, level: int) -> None:
            key = text.strip()
            if not key or key in seen_texts:
                return
            seen_texts.add(key)
            self.paragraphs.append({"text": key, "style": style, "level": level})

        for para in self.doc.paragraphs:
            style = para.style.name if para.style else "Normal"
            level = self._heading_level(style)
            _add(para.text, style, level)

        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    _add(cell.text, "Table", 0)

        log.debug("Extracted %d unique paragraphs", len(self.paragraphs))

    @staticmethod
    def _heading_level(style: str) -> int:
        """Handle multi-digit heading levels (Heading 1 … Heading 12+).  [Fix #16]"""
        m = re.match(r"Heading\s+(\d+)", style, re.IGNORECASE)
        if m:
            return int(m.group(1))
        if style.lower() in ("title", "subtitle"):
            return 1
        return 0

    def _check_not_scanned(self) -> None:
        if not self.paragraphs:
            _fatal(
                "Document appears empty or is image/scan-based. "
                "Clearfolio Review requires text-based .docx files."
            )
        snippet = " ".join(p["text"].lower() for p in self.paragraphs[:6])
        for kw in self.SCAN_KEYWORDS:
            if kw in snippet:
                _fatal(f"Scanned/image-based document detected ('{kw}').")

    def _build_full_text(self) -> None:
        """Build full text with hard cap to protect memory.  [Fix #07, #22]"""
        text = "\n".join(p["text"] for p in self.paragraphs)
        if len(text) > MAX_DOC_CHARS:
            self.truncated = True
            text = text[:MAX_DOC_CHARS]
            _warn(
                f"Document exceeds {MAX_DOC_CHARS:,} chars. "
                "Text truncated — later clauses may not be analysed."
            )
        self.full_text = text
        log.debug("Full text length: %d chars (truncated=%s)", len(self.full_text), self.truncated)


# ─────────────────────────────────────────────────────────────────────────────
# 2. DOCUMENT PARSER
# ─────────────────────────────────────────────────────────────────────────────
class DocumentParser:
    CLAUSE_START = re.compile(
        r"^(?:"
        r"(?:Article|Section|Clause|Schedule|Exhibit|Annex|Appendix)\s+[\w\d]+"
        r"|(?:\d+\.)+\d*\s+"
        r"|\([a-z]\)\s+"
        r"|[IVXLC]+\.\s+"
        r")",
        re.IGNORECASE,
    )

    def __init__(self, loader: DocumentLoader):
        self.paragraphs = loader.paragraphs

    def parse(self) -> list[dict]:
        sections: list[dict] = []
        current: Optional[dict] = None

        for para in self.paragraphs:
            text = para["text"]
            level = para["level"]
            is_heading = level > 0 or para["style"].lower() in ("title", "subtitle")
            is_clause = bool(self.CLAUSE_START.match(text))

            if is_heading or is_clause:
                if current:
                    sections.append(current)
                # NOTE: index placeholder — re-assigned below  [Fix #03]
                current = {"title": text, "body": "", "index": -1}
            else:
                if current is None:
                    current = {"title": "Preamble", "body": "", "index": -1}
                sep = "\n" if current["body"] else ""
                current["body"] += sep + text

        if current:
            sections.append(current)

        if len(sections) <= 1:
            sections = self._fallback_split()

        # ── Re-assign indices AFTER full list is built  [Fix #03] ───────────
        for i, s in enumerate(sections):
            s["index"] = i
            s["full_text"] = (s["title"] + " " + s["body"]).lower()

        log.debug("Parsed %d sections", len(sections))
        return sections

    def _fallback_split(self) -> list[dict]:
        CHUNK, chunks, buf = 4, [], []
        for i, para in enumerate(self.paragraphs):
            buf.append(para["text"])
            if len(buf) == CHUNK or i == len(self.paragraphs) - 1:
                chunks.append({
                    "title": buf[0][:80],
                    "body": "\n".join(buf[1:]),
                    "index": -1,  # will be fixed by caller
                    "full_text": " ".join(buf).lower(),
                })
                buf = []
        return chunks


# ─────────────────────────────────────────────────────────────────────────────
# 3. RULE ENGINE — 30 Deterministic Rules
# ─────────────────────────────────────────────────────────────────────────────
class RuleEngine:
    SEVERITY = {"CRITICAL": 4, "HIGH": 3, "MEDIUM": 2, "LOW": 1, "INFO": 0}

    # ── Precompiled document-level patterns  [Fix #08] ───────────────────────
    _DOC_PATTERNS: list[tuple] = [
        ("R01", "Missing Clause", "CRITICAL", "No Termination Clause",
         "No termination provision found. Parties may have no clear exit mechanism.",
         re.compile(r"\bterminat\w*\b")),
        ("R02", "Missing Clause", "CRITICAL", "No Liability Clause",
         "No clause limiting or defining liability. Unlimited exposure risk.",
         re.compile(r"\b(liabilit\w*|limitation of liability)\b")),
        ("R03", "Missing Clause", "HIGH", "No Jurisdiction / Governing Law",
         "Disputes may be subject to uncertain legal venue.",
         re.compile(r"\b(jurisdict\w*|governing law|choice of law|applicable law)\b")),
        ("R04", "Missing Clause", "HIGH", "No Payment Terms",
         "No payment, fee, or compensation clause detected.",
         re.compile(r"\b(payment|invoic\w*|fee|compensation|remunerat\w*)\b")),
        ("R05", "Missing Clause", "MEDIUM", "No Confidentiality Clause",
         "Sensitive information may not be protected.",
         re.compile(r"\b(confidential\w*|non-disclosure|nda|proprietary information)\b")),
        ("R06", "Missing Clause", "MEDIUM", "No Dispute Resolution Mechanism",
         "No arbitration, mediation, or ADR clause found.",
         re.compile(r"\b(arbitrat\w*|mediat\w*|dispute resolution|adr)\b")),
        ("R07", "Missing Clause", "MEDIUM", "No Force Majeure Clause",
         "Parties may be liable even under extraordinary circumstances.",
         re.compile(r"\b(force majeure|act of god|unforeseeable event)\b")),
        ("R08", "Missing Clause", "HIGH", "No Intellectual Property Clause",
         "Ownership of created works or IP may be ambiguous.",
         re.compile(r"\b(intellectual property|ip rights|copyright|ownership of work)\b")),
        ("R09", "Missing Clause", "MEDIUM", "No Indemnification Clause",
         "No indemnity provision — loss allocation is unclear.",
         re.compile(r"\bindemnif\w*\b")),
        ("R10", "Missing Clause", "LOW", "No Warranty or Disclaimer",
         "No warranty terms or disclaimer of warranties found.",
         re.compile(r"\b(warrant\w*|disclaim\w*|as-is|no representation)\b")),
        ("R11", "Missing Clause", "LOW", "No Amendment Procedure",
         "No process defined for modifying the agreement.",
         re.compile(r"\b(amend\w*|modif\w*|variation|change to this agreement)\b")),
        ("R12", "Missing Clause", "LOW", "No Notice Clause",
         "Method and address for official notices not specified.",
         re.compile(r"\b(notice|notification|written notice|notify)\b")),
        ("R13", "Missing Clause", "LOW", "No Assignment Clause",
         "Unclear whether rights can be transferred to third parties.",
         re.compile(r"\b(assign\w*|transfer of rights|delegate)\b")),
        ("R14", "Term Risk", "MEDIUM", "No Contract Term / Duration Defined",
         "Agreement may be open-ended with no expiry or renewal terms.",
         re.compile(r"\b(term|duration|period|expir\w*|renew\w*)\b")),
        ("R15", "Missing Clause", "LOW", "No Entire Agreement / Merger Clause",
         "Prior negotiations may override the contract.",
         re.compile(r"\b(entire agreement|whole agreement|supersedes|merger clause)\b")),
    ]

    # ── Precompiled clause-level patterns  [Fix #08] ─────────────────────────
    _CLAUSE_PATTERNS: list[tuple] = [
        ("R16", "One-sided Language", "HIGH", "Unilateral Termination / Sole Discretion",
         re.compile(r"\b(at\s+(its|our|their)\s+sole\s+discretion|without\s+cause|at\s+will)\b")),
        ("R17", "One-sided Language", "HIGH", "Unilateral Modification Right",
         re.compile(r"\b(reserves?\s+the\s+right\s+to\s+(change|modify|amend|update)|"
                    r"may\s+(change|modify|amend)\s+at\s+any\s+time)\b")),
        ("R18", "Liability Risk", "CRITICAL", "Unlimited Liability Exposure",
         re.compile(r"\b(unlimited\s+liability|no\s+cap\s+on\s+damages|"
                    r"liable\s+for\s+all\s+(losses|damages|costs))\b")),
        ("R20", "IP Risk", "HIGH", "Broad / Total IP Assignment",
         re.compile(r"\b(all\s+(rights|ip|intellectual property|inventions)\s+(are\s+)?assigned|"
                    r"work\s+made\s+for\s+hire|assign\s+all\s+rights?)\b")),
        ("R23", "Liability Risk", "MEDIUM", "Liquidated Damages Clause",
         re.compile(r"\bliquidated\s+damages?\b")),
        ("R24", "Payment Risk", "MEDIUM", "Penalties / Late Fees Defined",
         re.compile(r"\b(penalty|penalties|late\s+fee|interest\s+on\s+(late|unpaid))\b")),
        ("R25", "Legal Rights", "HIGH", "Jury Trial Waiver",
         re.compile(r"\bwaiv\w*\s+(jury\s+trial|right\s+to\s+jury)\b")),
        ("R26", "Legal Rights", "HIGH", "Class Action Waiver",
         re.compile(r"\bclass\s+action\s+waiv\w*|waiv\w*\s+class\s+action\b")),
        ("R29", "Payment Risk", "HIGH", "Unilateral Price Change Right",
         re.compile(r"\b(may\s+(increase|change|adjust)\s+(price|fee|rate|cost)|"
                    r"right\s+to\s+(change|adjust)\s+price)\b")),
    ]

    _RE_AUTO_RENEW = re.compile(r"\b(auto[- ]?renew\w*|automatically\s+renews?|rolls?\s+over)\b")
    _RE_OPT_OUT = re.compile(r"\b(opt[- ]out|cancel\w*|notice\s+to\s+(cancel|terminate|renew))\b")
    _RE_NON_COMPETE = re.compile(r"\b(non[- ]?compete|covenant\s+not\s+to\s+compete)\b")
    _RE_YEARS = re.compile(r"(\d+)\s+year")
    _RE_NON_SOLIC = re.compile(r"\bnon[- ]?solicit\w*\b")
    _RE_BEST_EFF = re.compile(r"\b(best\s+efforts?|reasonable\s+efforts?|commercially\s+reasonable)\b")
    _RE_SLA = re.compile(r"\b(\d+\s*(hour|day|week|month)|sla|service\s+level)\b")
    _RE_3P_INDEM = re.compile(r"\bindemnif\w+.{0,40}third[- ]?party\s+claim\b")
    _RE_SURVIVE_KEY = re.compile(r"\b(confidentialit\w*|intellectual property|indemnif\w*)\b")
    _RE_SURVIVE = re.compile(r"\bsurviv\w*\b")

    def __init__(self, clauses: list[dict], full_text: str):
        self.clauses = clauses
        self.full_text = full_text.lower()
        self._findings: list[dict] = []

    def run(self) -> list[dict]:
        self._findings = []
        self._document_level()
        self._clause_level()
        self._findings.sort(key=lambda f: -self.SEVERITY.get(f["severity"], 0))
        return self._findings

    def _document_level(self) -> None:
        ft = self.full_text
        for rule_id, cat, sev, title, detail, pattern in self._DOC_PATTERNS:
            if not pattern.search(ft):
                self._add(rule_id, cat, sev, title, detail, clause_index=-1)

    def _clause_level(self) -> None:
        for clause in self.clauses:
            ft = clause["full_text"]
            ci = clause["index"]

            # ── Fixed simple pattern checks ─────────────────────────────────
            for rule_id, cat, sev, title, pattern in self._CLAUSE_PATTERNS:
                if pattern.search(ft):
                    self._add(rule_id, cat, sev, title,
                              f"Clause '{clause['title'][:60]}' triggers this check.", ci)

            # ── R19 auto-renewal ─────────────────────────────────────────────
            if self._RE_AUTO_RENEW.search(ft) and not self._RE_OPT_OUT.search(ft):
                self._add("R19", "Term Risk", "MEDIUM",
                          "Auto-Renewal Without Opt-Out Mechanism",
                          f"Clause '{clause['title'][:60]}' auto-renews with no cancellation path.", ci)

            # ── R21 non-compete >1 year ──────────────────────────────────────
            if self._RE_NON_COMPETE.search(ft):
                years = self._RE_YEARS.findall(ft)
                if years and any(int(y) > 1 for y in years):
                    self._add("R21", "Restrictive Covenant", "HIGH",
                              "Non-Compete Exceeds 1 Year",
                              f"Clause '{clause['title'][:60]}' imposes non-compete > 1 year.", ci)

            # ── R22 non-solicitation ─────────────────────────────────────────
            if self._RE_NON_SOLIC.search(ft):
                self._add("R22", "Restrictive Covenant", "MEDIUM",
                          "Non-Solicitation Clause Present",
                          f"Clause '{clause['title'][:60]}' restricts solicitation.", ci)

            # ── R27 vague SLA ────────────────────────────────────────────────
            if self._RE_BEST_EFF.search(ft) and not self._RE_SLA.search(ft):
                self._add("R27", "Performance Risk", "LOW",
                          "Vague Performance Standard (Best Efforts Only)",
                          f"Clause '{clause['title'][:60]}' uses vague effort language with no SLA.", ci)

            # ── R28 broad third-party indemnity ──────────────────────────────
            if self._RE_3P_INDEM.search(ft):
                self._add("R28", "Indemnity Risk", "MEDIUM",
                          "Broad Third-Party Indemnification",
                          f"Clause '{clause['title'][:60]}' requires broad third-party indemnity.", ci)

        # ── R30 survival (document-level, fire once) ─────────────────────────
        if (self._RE_SURVIVE_KEY.search(self.full_text) and
                not self._RE_SURVIVE.search(self.full_text)):
            self._add("R30", "Survivability", "LOW",
                      "No Survival Clause for Key Obligations",
                      "Confidentiality / IP / indemnity may not survive termination.")

    def _add(self, rule_id: str, category: str, severity: str,
             title: str, detail: str, clause_index: int = -1) -> None:
        """
        Clause-level rules (clause_index >= 0) allow duplicate rule_id across
        different clauses.  Document-level rules (clause_index == -1) are
        global singletons.  [Fix #02]
        """
        if clause_index < 0:
            # Document-level: only one finding per rule_id
            if any(f["rule_id"] == rule_id for f in self._findings):
                return
        else:
            # Clause-level: one finding per (rule_id, clause_index) pair
            if any(f["rule_id"] == rule_id and f["clause_index"] == clause_index
                   for f in self._findings):
                return

        self._findings.append({
            "rule_id": rule_id,
            "category": category,
            "severity": severity,
            "title": title,
            "detail": detail,
            "clause_index": clause_index,
        })
        log.debug("Rule %s fired on clause_index=%d", rule_id, clause_index)


# ─────────────────────────────────────────────────────────────────────────────
# 4a. PROMPT CACHE  [Fix #19]
# ─────────────────────────────────────────────────────────────────────────────
class PromptCache:
    """
    In-process SHA-256 cache for Ollama prompt→response pairs.
    Updated: Now thread-safe for parallel model execution.
    """

    def __init__(self) -> None:
        self._store: dict[str, str] = {}
        self._lock = threading.Lock()

    def key(self, model: str, system: str, prompt: str) -> str:
        digest = hashlib.sha256(
            (model + "||" + system + "||" + prompt).encode()
        ).hexdigest()
        return digest

    def get(self, key: str) -> Optional[str]:
        with self._lock:
            return self._store.get(key)

    def set(self, key: str, value: str) -> None:
        with self._lock:
            self._store[key] = value

    def size(self) -> int:
        with self._lock:
            return len(self._store)


_prompt_cache = PromptCache()


# ─────────────────────────────────────────────────────────────────────────────
# 4b. OLLAMA CLIENT
# ─────────────────────────────────────────────────────────────────────────────
class OllamaClient:
    """
    Thin wrapper around Ollama /api/generate.
    Includes:
    - 3× retry with exponential back-off  [Fix #05]
    - tag-aware model resolution          [Fix #06, #23]
    - prompt cache integration            [Fix #19]
    - explicit model validation           [Fix #21]
    """

    def __init__(self, base_url: str = OLLAMA_BASE_URL,
                 timeout: int = OLLAMA_TIMEOUT) -> None:
        self.base_url = base_url
        self.timeout = timeout
        self._model_cache: Optional[list[str]] = None  # raw names with tags

    def is_alive(self) -> bool:
        try:
            r = requests.get(f"{self.base_url}/api/tags", timeout=5)
            return r.status_code == 200
        except Exception:
            return False

    def raw_models(self) -> list[str]:
        """Return full model names including :tag  [Fix #06]"""
        if self._model_cache is not None:
            return self._model_cache
        try:
            r = requests.get(f"{self.base_url}/api/tags", timeout=5)
            r.raise_for_status()
            self._model_cache = [m["name"] for m in r.json().get("models", [])]
        except Exception:
            self._model_cache = []
        log.debug("Available Ollama models: %s", self._model_cache)
        return self._model_cache

    def base_names(self) -> list[str]:
        """Return model names without :tag suffix"""
        return [m.split(":")[0] for m in self.raw_models()]

    def resolve_model(self, preferred: str, fallback: str) -> str:
        """
        Match preferred against available models handling :tag variants.
        Strategy:
          1. Exact match on full name (e.g. "llama3:8b")
          2. Base-name match  (e.g. "llama3" → "llama3:latest")
          3. Prefix match     (e.g. "deepseek-r1" → "deepseek-r1:7b")
          4. Fallback
        [Fix #06, #23]
        """
        raw = self.raw_models()
        bases = self.base_names()
        pref_base = preferred.split(":")[0]

        # 1. exact
        if preferred in raw:
            return preferred
        # 2. base-name
        if pref_base in bases:
            idx = bases.index(pref_base)
            log.debug("Resolved '%s' → '%s' (base match)", preferred, raw[idx])
            return raw[idx]
        # 3. prefix
        for full, base in zip(raw, bases):
            if base.startswith(pref_base):
                log.debug("Resolved '%s' → '%s' (prefix match)", preferred, full)
                return full

        # fallback
        if preferred != fallback:
            _warn(f"Model '{preferred}' not installed — trying fallback '{fallback}'")
        return self.resolve_model(fallback, fallback) if preferred != fallback else fallback

    def validate_model(self, model: str) -> None:
        """Abort with clear message if model is not resolvable.  [Fix #21]"""
        resolved = self.resolve_model(model, "")
        if not resolved:
            _fatal(
                f"Model '{model}' is not available and no fallback could be found. "
                f"Run: ollama pull {model}"
            )

    def generate(self, model: str, prompt: str, system: str,
                 temperature: float = 0.1) -> tuple[str, str]:
        """
        Call Ollama with retry + cache.
        Returns (response_text, error_string).  [Fix #05, #19]
        """
        cache_key = _prompt_cache.key(model, system, prompt)
        cached = _prompt_cache.get(cache_key)
        if cached is not None:
            log.debug("Cache hit for model=%s (key=%s…)", model, cache_key[:12])
            return cached, ""

        payload = {
            "model": model,
            "prompt": prompt,
            "system": system,
            "stream": False,
            "options": {"temperature": temperature, "num_predict": 1400},
        }
        last_err = ""
        for attempt in range(1, OLLAMA_RETRIES + 1):
            try:
                resp = requests.post(
                    f"{self.base_url}/api/generate",
                    json=payload,
                    timeout=self.timeout,
                )
                resp.raise_for_status()
                text = resp.json().get("response", "")
                _prompt_cache.set(cache_key, text)
                log.debug("Ollama ok model=%s attempt=%d", model, attempt)
                return text, ""
            except requests.exceptions.Timeout:
                last_err = "Request timed out"
            except Exception as exc:
                last_err = str(exc)
            if attempt < OLLAMA_RETRIES:
                delay = OLLAMA_RETRY_DELAY * (2 ** (attempt - 1))
                log.warning("Ollama attempt %d/%d failed (%s) — retry in %.1fs",
                            attempt, OLLAMA_RETRIES, last_err, delay)
                time.sleep(delay)

        log.error("Ollama failed after %d attempts: %s", OLLAMA_RETRIES, last_err)
        return "", last_err


# ─────────────────────────────────────────────────────────────────────────────
# 4c. RESPONSE PARSER
# ─────────────────────────────────────────────────────────────────────────────
VALID_RISKS = {"CRITICAL", "HIGH", "MEDIUM", "LOW"}


def parse_model_response(raw: str, model_name: str,
                         clauses: list[dict], pass_num: int = 1) -> list[AIFinding]:
    """
    Parse pipe-delimited rows.
    Clause index is matched by embedded [N] prefix in prompt; falls back to
    title fuzzy match.  [Fix #04, #10, #23]
    """
    if not raw:
        return []

    clause_map = {c["index"]: c for c in clauses}
    findings: list[AIFinding] = []

    for line in raw.strip().splitlines():
        line = line.strip()
        if not line or line.startswith("#") or set(line) <= set("-=+~"):
            continue

        # [Fix #10] — maxsplit=5 so extra pipes inside fields don't break parse
        parts = re.split(r"\s*\|\s*", line, maxsplit=5)
        if len(parts) < 5:
            continue

        clause_raw, risk_raw, issue, suggestion, conf_raw = parts[:5]
        risk = risk_raw.strip().upper()
        if risk not in VALID_RISKS:
            continue

        # ── Clause index resolution  [Fix #04] ──────────────────────────────
        matched_index = -1

        # Try embedded [N] prefix first (most reliable)
        idx_m = re.match(r"\[(\d+)]", clause_raw.strip())
        if idx_m:
            candidate = int(idx_m.group(1))
            if candidate in clause_map:
                matched_index = candidate
                clause_label = clause_raw[idx_m.end():].strip() or clause_map[candidate]["title"][:50]
            else:
                clause_label = clause_raw.strip()
        else:
            # Fuzzy title fallback
            clause_label = clause_raw.strip()
            label_lower = clause_label.lower()
            for idx, c in clause_map.items():
                t = c["title"].lower()
                if t.startswith(label_lower[:12]) or label_lower[:16] in t:
                    matched_index = idx
                    break

        # ── Confidence parse  [Fix #23] ─────────────────────────────────────
        conf_str = re.sub(r"[^\d.]", "", conf_raw)
        try:
            if not conf_str:
                raise ValueError("empty")
            cv = float(conf_str)
            conf_int = max(0, min(100, int(cv if cv > 1.0 else cv * 100)))
        except (ValueError, ZeroDivisionError):
            conf_int = 50

        findings.append(AIFinding(
            clause_label=clause_label[:50],
            risk=risk,
            issue=issue.strip()[:80],
            suggestion=suggestion.strip()[:80],
            confidence=conf_int,
            clause_index=matched_index,
            source_models=[model_name],
            pass_number=pass_num,
        ))

    log.debug("parse_model_response: %d findings from model=%s", len(findings), model_name)
    return findings


# ─────────────────────────────────────────────────────────────────────────────
# 4d. SYSTEM PROMPTS
# ─────────────────────────────────────────────────────────────────────────────
_TABLE_FORMAT_RULE = textwrap.dedent("""\
    Output ONLY a pipe-delimited table with NO header row and NO markdown:
        [N] Clause title | Risk | Issue | Suggestion | Confidence
    Rules:
    - [N]        : the integer index printed before each clause in the input, e.g. [3]
    - Clause title: first ≤8 words of the clause title
    - Risk       : exactly one of CRITICAL / HIGH / MEDIUM / LOW
    - Issue      : ≤15 words describing the risk
    - Suggestion : ≤15 words with a concrete fix
    - Confidence : integer percentage e.g. 78%
    Omit clauses that are well-drafted with no issues.
    Do NOT include any preamble, explanation, headers, or markdown fences.
""")

SYSTEM_PROMPT_A = textwrap.dedent("""\
    You are a CONTRACT RISK SPECIALIST. Your sole focus is identifying financial,
    legal, and operational risks inside contract clauses.
    Look for: uncapped liability, penalty exposure, unfavourable indemnity,
    one-sided termination, payment traps, auto-renewal without exit, price-change
    rights, and any clause that disproportionately burdens one party.
""") + _TABLE_FORMAT_RULE

SYSTEM_PROMPT_B = textwrap.dedent("""\
    You are a REGULATORY COMPLIANCE SPECIALIST. Your sole focus is ensuring
    contract clauses meet legal and regulatory requirements.
    Look for: missing mandatory disclosures, GDPR / data-protection obligations,
    consumer-protection violations, anti-competition clause issues, missing
    jurisdictional compliance, and unenforceable provisions.
""") + _TABLE_FORMAT_RULE

SYSTEM_PROMPT_C = textwrap.dedent("""\
    You are a CONTRACT COMPLETENESS SPECIALIST. Your sole focus is identifying
    what is MISSING or AMBIGUOUS in the contract.
    Look for: absent clauses that should exist (warranties, SLAs, IP ownership,
    force majeure, dispute resolution, confidentiality, survival), vague language
    lacking measurable definitions, undefined terms, and structural gaps that
    leave obligations unclear.
""") + _TABLE_FORMAT_RULE

SYSTEM_PROMPT_JUDGE = textwrap.dedent("""\
    You are the CHIEF CONTRACT AUDIT JUDGE with final authority.
    Input sections:
      [ORIGINAL CLAUSES] — source contract text indexed as [N] Title
      [MODEL A — Risk]         — findings from a risk specialist
      [MODEL B — Compliance]   — findings from a compliance specialist
      [MODEL C — Completeness] — findings from a completeness specialist

    Your job:
    1. DEDUPLICATE: merge findings about the same clause and same core issue
    2. RESOLVE CONFLICTS: when models disagree on Risk level, choose based on the original text
    3. VALIDATE: discard findings NOT supported by the original clause text
    4. ENRICH: keep the best Issue + Suggestion from whichever model phrased it best
    5. CONFIDENCE: boost if multiple models agree; reduce if only one and evidence is weak

    Output ONLY a pipe-delimited table, one row per validated finding, NO header:
        [N] Clause title | Risk | Issue | Suggestion | Confidence | Models | Note
    - [N]        : clause index integer (same as input)
    - Risk       : CRITICAL / HIGH / MEDIUM / LOW
    - Issue      : ≤15 words
    - Suggestion : ≤15 words
    - Confidence : integer percentage e.g. 84%
    - Models     : comma-separated source models e.g. A,B or A,B,C
    - Note       : ≤10 words judge rationale or "validated"
    No preamble, markdown, or extra text.
""")


# ─────────────────────────────────────────────────────────────────────────────
# 4e. CLAUSE PROMPT BUILDER  (shared)
# ─────────────────────────────────────────────────────────────────────────────
def chunk_clauses(clauses, size=4):
    return [clauses[i:i + size] for i in range(0, len(clauses), size)]


def build_clause_prompt(clauses: list[dict],
                        max_chars: int = MAX_CHARS_PER_CLAUSE) -> str:
    """
    Embed [index] before each clause so the model can return it verbatim.
    This enables exact index-based clause matching in parse_model_response.
    [Fix #04]
    """
    parts = []
    for c in clauses:
        snippet = (c["title"] + " " + c["body"])[:max_chars].replace("\n", " ")
        parts.append(f"[{c['index']}] {snippet}")
    return "\n\n".join(parts)


# ─────────────────────────────────────────────────────────────────────────────
# 4f. QUICK ANALYSER  (single model — quick mode)
# ─────────────────────────────────────────────────────────────────────────────
class QuickAnalyser:
    """Single-model pass.  [Fix #24 — returns AIFinding natively]"""

    def __init__(self, clauses: list[dict], client: OllamaClient) -> None:
        self.clauses = clauses
        self.client = client
        self.model = client.resolve_model(MODEL_QUICK, "llama3")

    def run(self) -> list[AIFinding]:
        if not self.client.is_alive():
            _banner_skipped("Ollama is not reachable at " + OLLAMA_BASE_URL)
            return []

        batch = self._select_clauses()
        prompt = build_clause_prompt(batch)

        _status(f"  Quick model [{self.model}] — analysing {len(batch)} clauses …")
        t0 = time.time()
        raw, err = self.client.generate(self.model, prompt, SYSTEM_PROMPT_A)
        elapsed = time.time() - t0

        if err:
            _banner_skipped(f"Quick model error after {elapsed:.1f}s: {err}")
            return []

        findings = parse_model_response(raw, self.model, self.clauses, pass_num=1)
        _status(f"  Quick model done in {elapsed:.1f}s — {len(findings)} findings")
        return findings

    def _select_clauses(self) -> list[dict]:
        return sorted(
            [c for c in self.clauses if len(c["body"]) > 30],
            key=lambda c: -len(c["body"])
        )[:QUICK_MAX_CLAUSES]


# ─────────────────────────────────────────────────────────────────────────────
# 4g. SPECIALIST RUNNER  (one role, one model)
# ─────────────────────────────────────────────────────────────────────────────
class SpecialistRunner:
    def __init__(self, model_name: str, role: str, system_prompt: str,
                 clauses: list[dict], client: OllamaClient,
                 pass_number: int = 1) -> None:
        self.model_name = model_name
        self.role = role
        self.system_prompt = system_prompt
        self.clauses = clauses
        self.client = client
        self.pass_number = pass_number

    def run(self) -> ModelResult:
        prompt = build_clause_prompt(self.clauses)
        t0 = time.time()
        raw, err = self.client.generate(self.model_name, prompt, self.system_prompt)
        elapsed = time.time() - t0

        if err:
            return ModelResult(model_name=self.model_name, role=self.role,
                               raw_output="", findings=[], elapsed_sec=elapsed,
                               error=err, pass_number=self.pass_number)

        findings = parse_model_response(raw, self.model_name,
                                        self.clauses, self.pass_number)
        return ModelResult(model_name=self.model_name, role=self.role,
                           raw_output=raw, findings=findings,
                           elapsed_sec=elapsed, pass_number=self.pass_number)


# ─────────────────────────────────────────────────────────────────────────────
# 4h. MULTI-MODEL PARALLEL LAYER
# ─────────────────────────────────────────────────────────────────────────────
class MultiModelLayer:
    """
    Run 3 specialist models in parallel.
    Uses per-future timeout instead of as_completed iterator timeout.  [Fix #17]
    """

    def __init__(self, clauses: list[dict], client: OllamaClient) -> None:
        self.clauses = clauses
        self.client = client
        self.model_a = client.resolve_model(MODEL_A, "llama3")
        self.model_b = client.resolve_model(MODEL_B, self.model_a)
        self.model_c = client.resolve_model(MODEL_C, self.model_a)
        log.debug("MultiModelLayer models: A=%s B=%s C=%s",
                  self.model_a, self.model_b, self.model_c)

    def run_pass(self, pass_number: int,
                 focus_clauses: Optional[list[dict]] = None) -> list[ModelResult]:
        if focus_clauses is not None:
            target = focus_clauses
        else:
            if DEEP_MAX_CLAUSES_PASS1 is None:
                target = self.clauses
            else:
                target = focus_clauses if focus_clauses is not None else self.clauses
        if not target:
            return []

        chunks = chunk_clauses(target)

        results: list[ModelResult] = []

        for chunk_id, chunk in enumerate(chunks):
            _status(f"  Chunk {chunk_id + 1}/{len(chunks)} — {len(chunk)} clauses")

            runners = [
                SpecialistRunner(self.model_a, "risk", SYSTEM_PROMPT_A,
                                 chunk, self.client, pass_number),
                SpecialistRunner(self.model_b, "compliance", SYSTEM_PROMPT_B,
                                 chunk, self.client, pass_number),
                SpecialistRunner(self.model_c, "completeness", SYSTEM_PROMPT_C,
                                 chunk, self.client, pass_number),
            ]

            with ThreadPoolExecutor(max_workers=3) as pool:
                futures = {pool.submit(r.run): r for r in runners}

                for fut, runner in futures.items():
                    try:
                        result = fut.result(timeout=OLLAMA_TIMEOUT + 30)
                        results.append(result)

                        icon = "✓" if not result.error else "✗"
                        _status(
                            f"    [{icon}] {runner.role:<13} [{runner.model_name}] "
                            f"{len(result.findings)} findings ({result.elapsed_sec:.1f}s)"
                        )

                    except Exception as exc:
                        _warn(f"{runner.model_name} ({runner.role}) failed: {exc}")

        return results

    def _select_clauses(self, n: int) -> list[dict]:
        return sorted(
            [c for c in self.clauses if len(c["body"]) > 30],
            key=lambda c: -len(c["body"])
        )[:n]


# ─────────────────────────────────────────────────────────────────────────────
# 4i. JUDGE MODEL
# ─────────────────────────────────────────────────────────────────────────────
class JudgeModel:
    CLAUSE_SNIPPET = 350

    def __init__(self, clauses: list[dict], client: OllamaClient) -> None:
        self.clauses = clauses
        self.client = client
        self.judge_model = client.resolve_model(JUDGE_MODEL, MODEL_A)

    def adjudicate(self, model_results: list[ModelResult]) -> list[AIFinding]:
        if not model_results:
            return []

        prompt = self._build_judge_prompt(model_results)
        _status(f"  Judge [{self.judge_model}] synthesising {len(model_results)} model outputs …")
        t0 = time.time()
        raw, err = self.client.generate(
            self.judge_model, prompt, SYSTEM_PROMPT_JUDGE, temperature=0.05
        )
        elapsed = time.time() - t0

        if err:
            _warn(f"Judge model error ({elapsed:.1f}s): {err} — using fallback merge")
            return self._fallback_merge(model_results)

        findings = self._parse_judge_response(raw)
        _status(f"  Judge done in {elapsed:.1f}s — {len(findings)} validated findings")
        return findings

    def _build_judge_prompt(self, model_results: list[ModelResult]) -> str:
        sections: list[str] = []

        # Original clauses with [N] index  [Fix #04]
        snippets = []
        for c in sorted(self.clauses, key=lambda x: -len(x["body"]))[:DEEP_MAX_CLAUSES_PASS1]:
            snippet = (c["title"] + " " + c["body"])[:self.CLAUSE_SNIPPET].replace("\n", " ")
            snippets.append(f"[{c['index']}] {snippet}")
        sections.append("[ORIGINAL CLAUSES]\n" + "\n\n".join(snippets))

        role_labels = {
            "risk": "A — Risk",
            "compliance": "B — Compliance",
            "completeness": "C — Completeness",
        }
        for result in model_results:
            if result.error or not result.raw_output.strip():
                continue
            label = role_labels.get(result.role, result.role.upper())
            header = (f"[MODEL {label} | model={result.model_name} "
                      f"| pass={result.pass_number}]")
            sections.append(header + "\n" + result.raw_output.strip())

        return "\n\n" + ("═" * 60 + "\n\n").join(sections)

    def _parse_judge_response(self, raw: str) -> list[AIFinding]:
        findings: list[AIFinding] = []
        clause_map = {c["index"]: c for c in self.clauses}

        for line in raw.strip().splitlines():
            line = line.strip()
            if not line or set(line) <= set("-=+~"):
                continue
            parts = re.split(r"\s*\|\s*", line, maxsplit=6)
            if len(parts) < 5:
                continue

            clause_raw = parts[0]
            risk = parts[1].strip().upper() if len(parts) > 1 else ""
            issue = parts[2].strip() if len(parts) > 2 else ""
            suggestion = parts[3].strip() if len(parts) > 3 else ""
            conf_raw = parts[4].strip() if len(parts) > 4 else "70"
            models_raw = parts[5].strip() if len(parts) > 5 else ""
            note = parts[6].strip() if len(parts) > 6 else "validated"

            if risk not in VALID_RISKS:
                continue

            # Index extraction  [Fix #04]
            matched_index = -1
            clause_label = clause_raw
            idx_m = re.match(r"\[(\d+)\]", clause_raw.strip())
            if idx_m:
                candidate = int(idx_m.group(1))
                if candidate in clause_map:
                    matched_index = candidate
                    clause_label = clause_raw[idx_m.end():].strip() or clause_map[candidate]["title"][:50]

            conf_str = re.sub(r"[^\d.]", "", conf_raw)
            try:
                cv = float(conf_str) if conf_str else 70.0
                conf_int = max(0, min(100, int(cv if cv > 1.0 else cv * 100)))
            except (ValueError, ZeroDivisionError):
                conf_int = 70

            source = [m.strip() for m in models_raw.split(",") if m.strip()] or ["judge"]
            findings.append(AIFinding(
                clause_label=clause_label[:50],
                risk=risk,
                issue=issue[:80],
                suggestion=suggestion[:80],
                confidence=conf_int,
                clause_index=matched_index,
                source_models=source,
                judge_validated=True,
                judge_note=note[:60],
                pass_number=0,
            ))

        return findings

    @staticmethod
    def _fallback_merge(model_results: list[ModelResult]) -> list[AIFinding]:
        """
        Programmatic dedup when judge is unavailable.
        Key includes issue prefix to prevent collisions.  [Fix #18]
        """
        RISK_RANK = {"CRITICAL": 4, "HIGH": 3, "MEDIUM": 2, "LOW": 1}
        bucket: dict[str, AIFinding] = {}

        for mr in model_results:
            for f in mr.findings:
                # Richer key: clause prefix + risk + issue prefix  [Fix #18]
                key = (
                    f.clause_label[:10].lower(),
                    f.risk,
                    f.issue[:12].lower(),
                )
                key_str = "|".join(key)
                if key_str not in bucket:
                    bucket[key_str] = f
                else:
                    existing = bucket[key_str]
                    if f.confidence > existing.confidence:
                        f.source_models = list(set(existing.source_models + f.source_models))
                        bucket[key_str] = f
                    else:
                        existing.source_models = list(set(existing.source_models + f.source_models))

        merged = sorted(
            bucket.values(),
            key=lambda f: (RISK_RANK.get(f.risk, 0), f.confidence),
            reverse=True,
        )
        for f in merged:
            f.judge_validated = False
            f.judge_note = "programmatic merge (judge unavailable)"
        return merged


# ─────────────────────────────────────────────────────────────────────────────
# 4j. DEEP AUDIT ORCHESTRATOR
# ─────────────────────────────────────────────────────────────────────────────
class DeepAuditOrchestrator:
    """
    Pass 1 — 3 models × top N clauses (breadth)
    Pass 2 — 3 models × high-risk clauses only (depth / confirmation)
    Judge  — synthesise → final findings
    Inter-pass sleep prevents Ollama queue saturation  [Fix #20]
    """

    def __init__(self, clauses: list[dict], client: OllamaClient) -> None:
        self.clauses = clauses
        self.client = client
        self.multi_layer = MultiModelLayer(clauses, client)
        self.judge = JudgeModel(clauses, client)
        self.all_results: list[ModelResult] = []
        self.final_findings: list[AIFinding] = []

    def run(self) -> list[AIFinding]:
        if not self.client.is_alive():
            _banner_skipped("Ollama is not reachable at " + OLLAMA_BASE_URL)
            return []

        # ── Pass 1 ────────────────────────────────────────────────────────────
        print()
        _status(col("Pass 1 / 2 — Broad analysis (3 specialist models in parallel)",
                    C.CYAN, C.BOLD))
        p1 = self.multi_layer.run_pass(pass_number=1)
        self.all_results.extend(p1)

        # ── Inter-pass pause  [Fix #20] ───────────────────────────────────────
        if INTER_PASS_SLEEP > 0:
            time.sleep(INTER_PASS_SLEEP)

        # ── Pass 2 ────────────────────────────────────────────────────────────
        focus = self._high_risk_clauses(p1)
        if focus:
            print()
            _status(col(f"Pass 2 / 2 — Deep re-analysis on {len(focus)} "
                        f"high-risk clause(s)", C.CYAN, C.BOLD))
            p2 = self.multi_layer.run_pass(pass_number=2, focus_clauses=focus)
            self.all_results.extend(p2)
            if INTER_PASS_SLEEP > 0:
                time.sleep(INTER_PASS_SLEEP)
        else:
            _status("  Pass 2 skipped — no CRITICAL/HIGH clauses in pass 1")

        # ── Judge ─────────────────────────────────────────────────────────────
        print()
        _status(col("Judge synthesis — deduplication + conflict resolution",
                    C.CYAN, C.BOLD))
        self.final_findings = self.judge.adjudicate(self.all_results)
        return self.final_findings

    def _high_risk_clauses(self, results: list[ModelResult]) -> list[dict]:
        idx_set: set[int] = set()
        for mr in results:
            for f in mr.findings:
                if f.risk in ("CRITICAL", "HIGH") and f.clause_index >= 0:
                    idx_set.add(f.clause_index)
        by_idx = {c["index"]: c for c in self.clauses}
        return [by_idx[i] for i in idx_set if i in by_idx][:DEEP_MAX_CLAUSES_PASS2]

    @property
    def model_results(self) -> list[ModelResult]:
        return self.all_results


# ─────────────────────────────────────────────────────────────────────────────
# 5. FINAL AGGREGATOR
# ─────────────────────────────────────────────────────────────────────────────
class FindingAggregator:
    RISK_RANK = {"CRITICAL": 4, "HIGH": 3, "MEDIUM": 2, "LOW": 1}

    @staticmethod
    def deduplicate_ai(findings: list[AIFinding]) -> list[AIFinding]:
        """Dedup by (clause_label_prefix, risk, issue_prefix).  [Fix #18]"""
        seen: dict[str, AIFinding] = {}
        for f in findings:
            key = (f"{f.clause_label[:10].lower()}|{f.risk}|{f.issue[:10].lower()}")
            if key not in seen or f.confidence > seen[key].confidence:
                seen[key] = f
        return sorted(
            seen.values(),
            key=lambda f: (FindingAggregator.RISK_RANK.get(f.risk, 0), f.confidence),
            reverse=True,
        )

    @staticmethod
    def merge_stats(rule_findings: list[dict],
                    ai_findings: list[AIFinding]) -> dict:
        counts = {"CRITICAL": 0, "HIGH": 0, "MEDIUM": 0, "LOW": 0}
        for f in rule_findings:
            counts[f.get("severity", "LOW")] = counts.get(f.get("severity", "LOW"), 0) + 1
        for f in ai_findings:
            counts[f.risk] = counts.get(f.risk, 0) + 1
        return counts


# ─────────────────────────────────────────────────────────────────────────────
# 6. AUDIT SCORER  [Fix #12]
# ─────────────────────────────────────────────────────────────────────────────
class AuditScorer:
    """
    Accumulates penalties separately before final subtraction.  [Fix #12]
    """
    RULE_PENALTY = {"CRITICAL": 12, "HIGH": 7, "MEDIUM": 4, "LOW": 2, "INFO": 0}
    AI_PENALTY = {"CRITICAL": 6, "HIGH": 4, "MEDIUM": 2, "LOW": 1}

    def compute(self, rule_findings: list[dict],
                ai_findings: list[AIFinding]) -> int:
        rule_total = sum(
            self.RULE_PENALTY.get(f["severity"], 0) for f in rule_findings
        )
        ai_total = sum(
            self.AI_PENALTY.get(f.risk, 0) for f in ai_findings
        )
        total_penalty = rule_total + ai_total
        score = 100 - total_penalty
        log.debug("Score: 100 - %d(rule) - %d(ai) = %d",
                  rule_total, ai_total, score)
        return max(0, min(100, score))

    @staticmethod
    def grade(score: int) -> tuple[str, str]:
        if score >= 85:
            return "A", C.GREEN
        if score >= 70:
            return "B", C.CYAN
        if score >= 55:
            return "C", C.YELLOW
        if score >= 40:
            return "D", C.MAGENTA
        return "F", C.RED


# ─────────────────────────────────────────────────────────────────────────────
# 7. OUTPUT GENERATOR — terminal report
# ─────────────────────────────────────────────────────────────────────────────
class OutputGenerator:
    SEV_COLOUR = {
        "CRITICAL": C.RED,
        "HIGH": C.MAGENTA,
        "MEDIUM": C.YELLOW,
        "LOW": C.CYAN,
        "INFO": C.DIM,
    }
    MODEL_COLOUR = {
        "risk": C.RED,
        "compliance": C.CYAN,
        "completeness": C.PURPLE,
    }

    def __init__(self, filepath: str, clauses: list[dict],
                 rule_findings: list[dict], ai_findings: list[AIFinding],
                 score: int, mode: str, ai_skipped: bool,
                 model_results: Optional[list[ModelResult]] = None) -> None:
        self.filepath = filepath
        self.clauses = clauses
        self.rule_findings = rule_findings
        self.ai_findings = ai_findings
        self.score = score
        self.mode = mode
        self.ai_skipped = ai_skipped
        self.model_results = model_results or []

    def render(self) -> None:
        self._header()
        self._summary_box()
        if self.mode == "deep" and self.model_results:
            self._model_panel()
        self._rule_findings()
        self._ai_findings()
        self._missing_clauses()
        if self.mode == "deep":
            self._judge_section()
        self._footer()

    # ── sections ─────────────────────────────────────────────────────────────
    def _header(self) -> None:
        print()
        print(col("╔" + "═" * (TERMINAL_WIDTH - 2) + "╗", C.CYAN, C.BOLD))
        title = "  CLEARFOLIO REVIEW v3 — CONTRACT AUDIT REPORT  "
        pad = (TERMINAL_WIDTH - 2 - len(title)) // 2
        print(col(
            "║" + " " * pad + title +
            " " * (TERMINAL_WIDTH - 2 - pad - len(title)) + "║",
            C.CYAN, C.BOLD
        ))
        print(col("╚" + "═" * (TERMINAL_WIDTH - 2) + "╝", C.CYAN, C.BOLD))
        print(col(f"  File    : {os.path.basename(self.filepath)}", C.WHITE))
        print(col(f"  Path    : {self.filepath}", C.DIM))
        print(col(f"  Audited : {datetime.now().strftime('%Y-%m-%d  %H:%M:%S')}", C.DIM))
        print(col(f"  Clauses : {len(self.clauses)}", C.DIM))
        print(col(f"  Mode    : {self.mode.upper()}", C.DIM))
        print()

    def _summary_box(self) -> None:
        grade, gcol = AuditScorer.grade(self.score)
        stats = FindingAggregator.merge_stats(self.rule_findings, self.ai_findings)
        bar_fill = int(self.score / 100 * 44)
        bar = col("█" * bar_fill, gcol) + col("░" * (44 - bar_fill), C.DIM)

        print(hr("═", C.BOLD))
        print(col("  AUDIT SCORE", C.BOLD, C.WHITE))
        print(hr("─"))
        print(f"  Score : {col(str(self.score) + ' / 100', gcol, C.BOLD)}"
              f"   Grade : {col(grade, gcol, C.BOLD)}")
        print(f"  [{bar}]")
        print()
        print(f"  Totals :  "
              f"{col(str(stats['CRITICAL']) + ' CRITICAL', C.RED, C.BOLD)}  "
              f"{col(str(stats['HIGH']) + ' HIGH', C.MAGENTA)}  "
              f"{col(str(stats['MEDIUM']) + ' MEDIUM', C.YELLOW)}  "
              f"{col(str(stats['LOW']) + ' LOW', C.CYAN)}")
        print(f"  Sources:  "
              f"Rule engine {col(str(len(self.rule_findings)), C.WHITE)}  "
              f"AI {col(str(len(self.ai_findings)), C.WHITE)}"
              + (col("  [AI skipped]", C.YELLOW) if self.ai_skipped else ""))
        print(hr("═", C.BOLD))
        print()

    def _model_panel(self) -> None:
        print(col("  MODEL PERFORMANCE SUMMARY", C.BOLD, C.WHITE))
        print(hr("─"))
        passes: dict[int, list[ModelResult]] = {}
        for mr in self.model_results:
            passes.setdefault(mr.pass_number, []).append(mr)
        for pass_num in sorted(passes.keys()):
            print(col(f"  Pass {pass_num}:", C.BOLD, C.CYAN))
            for mr in passes[pass_num]:
                rc = self.MODEL_COLOUR.get(mr.role, C.WHITE)
                ok = col("✓", C.GREEN) if not mr.error else col("✗", C.RED)
                print(f"    {ok}  {col(mr.model_name, rc):<28}  "
                      f"{col('[' + mr.role + ']', rc, C.DIM):<24}  "
                      f"{col(str(len(mr.findings)) + ' findings', C.WHITE):<16}  "
                      f"{col(str(round(mr.elapsed_sec, 1)) + 's', C.DIM)}"
                      + (col(f"  {mr.error[:40]}", C.RED) if mr.error else ""))
        print()

    def _rule_findings(self) -> None:
        if not self.rule_findings:
            print(col("  ✓ No deterministic rule violations found.", C.GREEN, C.BOLD))
            print()
            return
        print(col(f"  RULE ENGINE FINDINGS  ({len(self.rule_findings)})", C.BOLD, C.WHITE))
        print(hr("─"))
        for f in self.rule_findings:
            sev = f["severity"]
            sc = self.SEV_COLOUR.get(sev, C.WHITE)
            print(f"  {col('[' + sev + ']', sc, C.BOLD):<30} "
                  f"{col(f['rule_id'], C.DIM)}  "
                  f"{col(f['title'], C.WHITE, C.BOLD)}")
            wrapped = textwrap.fill(
                f["detail"], width=TERMINAL_WIDTH - 8,
                initial_indent=" " * 6, subsequent_indent=" " * 6
            )
            print(col(wrapped, C.DIM))
            if f["clause_index"] >= 0:
                ci = f["clause_index"]
                label = self.clauses[ci]["title"][:60] if ci < len(self.clauses) else "—"
                print(col(f"       ↳ Clause [{ci}]: {label}", C.DIM))
            print()
        print()

    def _ai_findings(self) -> None:
        if not self.ai_findings:
            msg = ("AI ANALYSIS SKIPPED" if self.ai_skipped
                   else "AI analysis returned no findings.")
            print(col(f"  ⊘  {msg}", C.YELLOW))
            print()
            return
        mode_label = ("MULTI-MODEL AI — JUDGE VALIDATED"
                      if self.mode == "deep" else f"AI ANALYSIS  [{MODEL_QUICK}]")
        print(col(f"  {mode_label}", C.BOLD, C.WHITE))
        print(hr("─"))
        cw = [26, 10, 27, 27, 7]
        print("  " + "  ".join([
                                   col(f"{'Clause':<{cw[0]}}", C.BOLD),
                                   col(f"{'Risk':<{cw[1]}}", C.BOLD),
                                   col(f"{'Issue':<{cw[2]}}", C.BOLD),
                                   col(f"{'Suggestion':<{cw[3]}}", C.BOLD),
                                   col(f"{'Conf':<{cw[4]}}", C.BOLD),
                               ] + ([col("Models", C.BOLD)] if self.mode == "deep" else [])))
        print("  " + "─" * (sum(cw) + 2 * (len(cw) + 1)))
        for f in self.ai_findings:
            rc = self.SEV_COLOUR.get(f.risk, C.WHITE)
            cc = C.GREEN if f.confidence >= 80 else (C.YELLOW if f.confidence >= 60 else C.RED)
            val_tag = col("✓", C.GREEN) if f.judge_validated else col("·", C.DIM)
            row = (
                    col(f"{f.clause_label[:cw[0]]:<{cw[0]}}", C.WHITE) + "  " +
                    col(f"{f.risk:<{cw[1]}}", rc, C.BOLD) + "  " +
                    f"{f.issue[:cw[2]]:<{cw[2]}}  " +
                    f"{f.suggestion[:cw[3]]:<{cw[3]}}  " +
                    col(f"{str(f.confidence) + '%':<{cw[4]}}", cc)
            )
            if self.mode == "deep":
                src = ",".join(f.source_models)[:10]
                row += "  " + col(f"{val_tag} {src}", C.DIM)
            print("  " + row)
        print()

    def _missing_clauses(self) -> None:
        ids = {f"R{str(i).zfill(2)}" for i in range(1, 16)}
        missing = [f for f in self.rule_findings if f["rule_id"] in ids]
        if not missing:
            print(col("  ✓ All standard clauses present.", C.GREEN))
            print()
            return
        print(col(f"  MISSING / ABSENT CLAUSES  ({len(missing)})", C.BOLD, C.RED))
        print(hr("─"))
        for f in missing:
            sc = self.SEV_COLOUR.get(f["severity"], C.WHITE)
            print(f"  {col('✗', sc, C.BOLD)}  {col(f['title'], C.WHITE)}"
                  f"  {col('[' + f['severity'] + ']', sc)}")
        print()

    def _judge_section(self) -> None:
        validated = [f for f in self.ai_findings if f.judge_validated]
        consensus = [f for f in self.ai_findings if len(f.source_models) >= 2]
        print(col("  JUDGE VALIDATION SUMMARY", C.BOLD, C.WHITE))
        print(hr("─"))
        print(f"  Judge-validated : {col(str(len(validated)), C.GREEN, C.BOLD)}")
        print(f"  Multi-model consensus (≥2 models) : {col(str(len(consensus)), C.CYAN)}")
        if consensus:
            print()
            for f in sorted(
                    consensus,
                    key=lambda x: ({"CRITICAL": 4, "HIGH": 3, "MEDIUM": 2, "LOW": 1}.get(x.risk, 0),
                                   x.confidence),
                    reverse=True,
            )[:6]:
                rc = self.SEV_COLOUR.get(f.risk, C.WHITE)
                src = ",".join(f.source_models)
                print(f"    {col(f.risk, rc, C.BOLD):<22}  "
                      f"{col(f.clause_label[:28], C.WHITE):<30}  "
                      f"src={col(src, C.DIM):<10}  "
                      f"conf={col(str(f.confidence) + '%', C.GREEN)}"
                      + (f"  {col(f.judge_note[:40], C.DIM)}" if f.judge_note else ""))
        print()

    def _footer(self) -> None:
        print(hr("═", C.BOLD))
        print(col("  Clearfolio Review v3 | Local-Only | No Data Leaves This Machine",
                  C.DIM))
        print(col("  " + "─" * (TERMINAL_WIDTH - 4), C.DIM))
        print(col("  DISCLAIMER: AI-assisted output. Does not constitute legal advice.",
                  C.DIM))
        print(hr("═", C.BOLD))
        print()


# ─────────────────────────────────────────────────────────────────────────────
# 8. REPORT EXPORT  [Fix #14]
# ─────────────────────────────────────────────────────────────────────────────
def export_report(filepath: str, clauses: list[dict],
                  rule_findings: list[dict], ai_findings: list[AIFinding],
                  score: int, grade: str, mode: str,
                  ai_skipped: bool, model_results: list[ModelResult],
                  fmt: str) -> str:
    """
    Write structured report to <input_basename>_audit.<ext>.
    Returns the output path.
    """
    base = os.path.splitext(filepath)[0]
    out_path = f"{base}_audit.{fmt}"

    if fmt == "json":
        data = {
            "clearfolio_version": "3",
            "audited_at": datetime.now().isoformat(),
            "filename": os.path.basename(filepath),
            "mode": mode,
            "score": score,
            "grade": grade,
            "clauses_count": len(clauses),
            "ai_skipped": ai_skipped,
            "rule_findings": rule_findings,
            "ai_findings": [asdict(f) for f in ai_findings],
            "model_results": [
                {
                    "model": mr.model_name,
                    "role": mr.role,
                    "pass": mr.pass_number,
                    "findings": len(mr.findings),
                    "elapsed_sec": round(mr.elapsed_sec, 1),
                    "error": mr.error or None,
                }
                for mr in model_results
            ],
        }
        with open(out_path, "w", encoding="utf-8") as fh:
            json.dump(data, fh, indent=2, default=str)

    elif fmt == "txt":
        lines: list[str] = [
            "CLEARFOLIO REVIEW v3 — AUDIT REPORT",
            "=" * 60,
            f"File    : {os.path.basename(filepath)}",
            f"Audited : {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            f"Mode    : {mode.upper()}",
            f"Score   : {score}/100  Grade: {grade}",
            f"Clauses : {len(clauses)}",
            "",
            "RULE FINDINGS",
            "-" * 40,
        ]
        for f in rule_findings:
            lines.append(f"[{f['severity']}] {f['rule_id']} {f['title']}")
            lines.append(f"    {f['detail']}")
            lines.append("")
        lines += ["", "AI FINDINGS", "-" * 40]
        for f in ai_findings:
            lines.append(f"[{f.risk}] {f.clause_label}  conf={f.confidence}%")
            lines.append(f"    Issue      : {f.issue}")
            lines.append(f"    Suggestion : {f.suggestion}")
            lines.append("")
        with open(out_path, "w", encoding="utf-8") as fh:
            fh.write("\n".join(lines))

    return out_path


# ─────────────────────────────────────────────────────────────────────────────
# CLI ARGUMENT PARSER
# ─────────────────────────────────────────────────────────────────────────────
def parse_args() -> argparse.Namespace:
    p = argparse.ArgumentParser(
        prog="clearfolio_review",
        description="Clearfolio Review v3 — Production-grade privacy-first contract auditor",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=textwrap.dedent("""\
            examples:
              python clearfolio_review.py nda.docx
              python clearfolio_review.py contract.docx --mode deep --output json
              CF_MODEL_A=llama3 CF_JUDGE_MODEL=deepseek-r1 \\
                  python clearfolio_review.py contract.docx --mode deep --verbose

            environment variables (all optional):
              CF_OLLAMA_URL    Ollama base URL         (default: http://localhost:11434)
              CF_TIMEOUT       Per-request timeout (s) (default: 180)
              CF_RETRIES       Retry attempts          (default: 3)
              CF_RETRY_DELAY   Retry delay (s)         (default: 2.0)
              CF_PASS_SLEEP    Inter-pass sleep (s)    (default: 1.5)
              CF_MODEL_QUICK   Quick mode model        (default: llama3)
              CF_MODEL_A       Risk specialist         (default: llama3)
              CF_MODEL_B       Compliance specialist   (default: mistral)
              CF_MODEL_C       Completeness specialist (default: phi3)
              CF_JUDGE_MODEL   Judge model             (default: deepseek-r1)
        """),
    )
    p.add_argument("filepath", help=".docx contract file path")
    p.add_argument(
        "--mode", choices=["quick", "deep"], default="quick",
        help="quick = single model (default) | deep = 3 models + judge",
    )
    p.add_argument(
        "--output", choices=["json", "txt"], default=None,
        help="write structured report file alongside input",
    )
    p.add_argument("--verbose", action="store_true",
                   help="enable DEBUG logging to stderr")
    p.add_argument("--info", action="store_true",
                   help="print architecture notes and exit")
    return p.parse_args()


def _print_info() -> None:
    """Fix #25 — architecture notes accessible via --info"""
    notes = textwrap.dedent("""\
    ╔══════════════════════════════════════════════════════════════╗
    ║    CLEARFOLIO REVIEW v3 — ARCHITECTURE NOTES                 ║
    ╚══════════════════════════════════════════════════════════════╝

    Pipeline:
      .docx → DocumentLoader → DocumentParser → RuleEngine → [AI] → Score → Output

    DocumentLoader
      Reads paragraphs + table cells (deduplicated).
      Caps full_text at 200 k chars to protect memory.
      Rejects scanned/empty documents.

    DocumentParser
      Detects clause boundaries via heading styles and numbered patterns.
      Assigns indices AFTER the full list is built (no off-by-one bugs).
      Falls back to paragraph-chunk grouping for unstructured documents.

    RuleEngine  (30 deterministic rules)
      R01–R15  document-level (absent clause checks) — one finding per rule.
      R16–R30  clause-level   (risky pattern checks) — one finding per
               (rule_id, clause_index) pair, so the same rule can fire in
               multiple clauses.
      All patterns are precompiled at class definition.

    OllamaClient
      All calls retry up to CF_RETRIES times with exponential back-off.
      Prompt/response pairs are SHA-256 cached in-process.
      Model resolution handles :tag suffixes on both preferred and available.

    AI Modes
      quick  — QuickAnalyser sends top 12 clauses to a single model.
      deep   — DeepAuditOrchestrator runs:
               Pass 1: 3 specialist models in parallel (risk / compliance / completeness)
               Pass 2: same 3 models on CRITICAL/HIGH clauses only
               Judge:  synthesises all outputs (dedup, conflict resolution, validation)

    parse_model_response
      Clauses are indexed [N] in prompts; models echo the index back.
      Parser uses exact index matching first, title-fuzzy as fallback.
      Uses re.split(maxsplit=5) — immune to extra pipe characters.

    Output
      Terminal (always) + optional --output json|txt report file.
    """)
    print(notes)


# ─────────────────────────────────────────────────────────────────────────────
# MAIN PIPELINE
# ─────────────────────────────────────────────────────────────────────────────
def main() -> None:
    args = parse_args()
    _configure_logging(args.verbose)

    if args.info:
        _print_info()
        sys.exit(0)

    print()
    mode_tag = col(f"[{args.mode.upper()} MODE]", C.BOLD,
                   C.GREEN if args.mode == "deep" else C.BLUE)
    print(col("  ■ CLEARFOLIO REVIEW v3 ", C.CYAN, C.BOLD) + mode_tag)
    print()

    # ── Step 1: Load ─────────────────────────────────────────────────────────
    _status("Loading document …")
    loader = DocumentLoader(args.filepath).load()
    _status(f"Loaded {len(loader.paragraphs)} paragraphs "
            f"from {col(os.path.basename(args.filepath), C.WHITE)}"
            + (col("  [truncated]", C.YELLOW) if loader.truncated else ""))

    # ── Step 2: Parse ─────────────────────────────────────────────────────────
    _status("Parsing clauses …")
    clauses = DocumentParser(loader).parse()
    _status(f"Detected {col(str(len(clauses)), C.WHITE)} clauses / sections")

    # ── Step 3: Rule Engine ───────────────────────────────────────────────────
    _status("Running rule engine (30 rules) …")
    rule_findings = RuleEngine(clauses, loader.full_text).run()
    _status(f"Rule engine: {col(str(len(rule_findings)), C.WHITE)} findings")

    # ── Step 4: AI Analysis ───────────────────────────────────────────────────
    client = OllamaClient()
    ai_findings: list[AIFinding] = []
    model_results: list[ModelResult] = []
    ai_skipped = False

    if args.mode == "quick":
        _section("Quick AI Analysis")
        ai_findings = QuickAnalyser(clauses, client).run()
        if not ai_findings and not client.is_alive():
            ai_skipped = True

    else:
        _section("Deep AI Analysis — Multi-Model + Judge Pipeline")
        orch = DeepAuditOrchestrator(clauses, client)
        ai_findings = orch.run()
        model_results = orch.model_results
        if not ai_findings and not client.is_alive():
            ai_skipped = True

    # ── Step 5: Deduplicate ───────────────────────────────────────────────────
    ai_findings = FindingAggregator.deduplicate_ai(ai_findings)
    _status(f"Final AI findings: {col(str(len(ai_findings)), C.WHITE)}"
            + (col("  (cache hits: " + str(_prompt_cache.size()) + ")", C.DIM)
               if _prompt_cache.size() else ""))

    # ── Step 6: Score ─────────────────────────────────────────────────────────
    score = AuditScorer().compute(rule_findings, ai_findings)
    grade, _gcol = AuditScorer.grade(score)

    # ── Step 7: Terminal report ────────────────────────────────────────────────
    _status("Generating report …\n")
    OutputGenerator(
        filepath=args.filepath,
        clauses=clauses,
        rule_findings=rule_findings,
        ai_findings=ai_findings,
        score=score,
        mode=args.mode,
        ai_skipped=ai_skipped,
        model_results=model_results,
    ).render()

    # ── Step 8: Optional file export  [Fix #14] ──────────────────────────────
    if args.output:
        out_path = export_report(
            filepath=args.filepath,
            clauses=clauses,
            rule_findings=rule_findings,
            ai_findings=ai_findings,
            score=score,
            grade=grade,
            mode=args.mode,
            ai_skipped=ai_skipped,
            model_results=model_results,
            fmt=args.output,
        )
        _status(col(f"Report saved → {out_path}", C.GREEN))


if __name__ == "__main__":
    main()
